import io
import PyPDF2
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


def word_to_pdf(docx_file):
    """Convert a .docx file-like object to PDF. Returns BytesIO."""
    doc = Document(docx_file)
    output = io.BytesIO()
    pdf = SimpleDocTemplate(
        output, pagesize=letter,
        rightMargin=50, leftMargin=50,
        topMargin=60, bottomMargin=60
    )
    styles = getSampleStyleSheet()
    story = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 8))
            continue

        style_name = para.style.name if para.style else ''
        if style_name.startswith('Heading 1'):
            style = styles['Heading1']
        elif style_name.startswith('Heading 2'):
            style = styles['Heading2']
        elif style_name.startswith('Heading 3'):
            style = styles['Heading3']
        else:
            style = styles['Normal']

        full_text = ''
        for run in para.runs:
            run_text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if run.bold and run.italic:
                full_text += f'<b><i>{run_text}</i></b>'
            elif run.bold:
                full_text += f'<b>{run_text}</b>'
            elif run.italic:
                full_text += f'<i>{run_text}</i>'
            else:
                full_text += run_text

        if full_text.strip():
            try:
                story.append(Paragraph(full_text, style))
                story.append(Spacer(1, 4))
            except Exception:
                clean = para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(clean, styles['Normal']))
                story.append(Spacer(1, 4))

    if not story:
        story.append(Paragraph('Empty document', styles['Normal']))

    pdf.build(story)
    output.seek(0)
    return output

def pdf_to_word(pdf_file):
    """Convert a PDF file-like object to .docx. Returns BytesIO."""
    reader = PyPDF2.PdfReader(pdf_file)
    doc = Document()
    doc.add_heading('Converted Document', level=1)

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            doc.add_heading(f'Page {i + 1}', level=2)
            for line in text.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        doc.add_page_break()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def images_to_pdf(image_files):
    """Convert multiple image files to a single PDF. Returns BytesIO."""
    from PIL import Image
    import io

    output = io.BytesIO()
    image_list = []

    for img_file in image_files:
        img = Image.open(img_file)
        if img.mode in ('RGBA', 'P', 'LA'):
            img = img.convert('RGB')
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        image_list.append(img)

    if not image_list:
        raise ValueError("No valid images provided.")

    first = image_list[0]
    rest  = image_list[1:]
    first.save(output, format='PDF', save_all=True, append_images=rest)
    output.seek(0)
    return output